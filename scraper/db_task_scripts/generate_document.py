from docx import Document
from docx.shared import Inches
import psycopg2
import pandas as pd
import re

conn = psycopg2.connect("dbname=polovniautomobili user=postgres password=123")
cur = conn.cursor()

# Return only one line
def read_script_from_file(script_location):
    with open(script_location, 'r', encoding='utf8') as script:
        temp = script.readlines()
        return "".join(map(str, temp))

def read_script_and_execute(script_location):
    sql_string = read_script_from_file(script_location)
    cur.execute(sql_string)
    headings = [desc[0] for desc in cur.description]
    return cur.fetchall(), headings


biggest_engine_script_location = './scripts/biggest_engine.sql'
biggest_strength_script_location = './scripts/biggest_strength.sql'
first_30_most_expensive_script_location = './scripts/first_30_most_expensive.sql'
first_30_most_expensive_suv_script_location = './scripts/first_30_most_expensive_suv.sql'
list_number_of_cars_depending_on_color_script_location = './scripts/list_number_of_cars_depending_on_color.sql'
list_number_of_cars_manufacturer_script_location = './scripts/list_number_of_cars_for_all_manufacturers.sql'
list_number_of_cars_per_city_script_location = './scripts/list_number_of_cars_per_city.sql'
longest_dist_script_location = './scripts/longest_dist.sql'
ranking_list_2021_2022_script_location = './scripts/ranking_list_2021_2022.sql'

biggest_engine_result, biggest_engine_result_headings = read_script_and_execute(biggest_engine_script_location)
biggest_strength_result, biggest_strength_result_headings = read_script_and_execute(biggest_strength_script_location)
first_30_most_expensive_result, first_30_most_expensive_result_headings = read_script_and_execute(first_30_most_expensive_script_location)
first_30_most_expensive_suv_result, first_30_most_expensive_suv_result_headings = read_script_and_execute(first_30_most_expensive_suv_script_location)
number_of_cars_color_result_result, number_of_cars_color_result_result_headings = read_script_and_execute(list_number_of_cars_depending_on_color_script_location)
number_of_cars_manufacturer_result_result, number_of_cars_manufacturer_result_result_headings = read_script_and_execute(list_number_of_cars_depending_on_color_script_location)
number_of_cars_per_city_result, number_of_cars_per_city_result_headings = read_script_and_execute(list_number_of_cars_per_city_script_location)
longest_dist_result, longest_dist_result_headings = read_script_and_execute(longest_dist_script_location)
ranking_list_result, ranking_list_result_headings = read_script_and_execute(ranking_list_2021_2022_script_location)

cur.close()
conn.close()

def populate_document_with_query_data(document, query_data, headings):
    number_of_rows = len(query_data)
    if number_of_rows == 0:
        print('Error')
        return
    number_of_cols = len(headings)

    table = document.add_table(rows=1, cols=number_of_cols, style='Light Grid Accent 1')
    hdr_cells = table.rows[0].cells
    for i, cell in enumerate(headings):
        hdr_cells[i].text = cell

    for row in query_data:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)

    document.add_page_break()



document = Document()

document.add_heading('Rezultati upita', 0)
document.add_page_break()

print('Zadatak a')
document.add_heading('Rezultati upita za zadatak pod a', level=1)
document.add_paragraph(read_script_from_file(list_number_of_cars_manufacturer_script_location))
populate_document_with_query_data(document, number_of_cars_manufacturer_result_result, number_of_cars_manufacturer_result_result_headings)

print('Zadatak b')
document.add_heading('Rezultati upita za zadatak pod b', level=1)
document.add_paragraph(read_script_from_file(list_number_of_cars_per_city_script_location))
populate_document_with_query_data(document, number_of_cars_per_city_result, number_of_cars_per_city_result_headings)

print('Zadatak c')
document.add_heading('Rezultati upita za zadatak pod c', level=1)
document.add_paragraph(read_script_from_file(list_number_of_cars_depending_on_color_script_location))
populate_document_with_query_data(document, number_of_cars_color_result_result, number_of_cars_color_result_result_headings)

print('Zadatak d.1.')
document.add_heading('Rezultati upita za zadatak pod d (obicno)', level=1)
document.add_paragraph(read_script_from_file(first_30_most_expensive_script_location))
populate_document_with_query_data(document, first_30_most_expensive_result, first_30_most_expensive_result_headings)

print('Zadatak d.2.')
document.add_heading('Rezultati upita za zadatak pod d (SUV)', level=1)
document.add_paragraph(read_script_from_file(first_30_most_expensive_suv_script_location))
populate_document_with_query_data(document, first_30_most_expensive_suv_result, first_30_most_expensive_suv_result_headings)

print('Zadatak e')
document.add_heading('Rezultati upita za zadatak pod e', level=1)
document.add_paragraph(read_script_from_file(ranking_list_2021_2022_script_location))
populate_document_with_query_data(document, ranking_list_result, ranking_list_result_headings)

print('Zadatak f')
document.add_heading('Rezultati upita za zadatak pod f', level=1)
document.add_paragraph(read_script_from_file(biggest_engine_script_location))
populate_document_with_query_data(document, biggest_engine_result, biggest_engine_result_headings)
document.add_paragraph(read_script_from_file(biggest_strength_script_location))
populate_document_with_query_data(document, biggest_strength_result, biggest_strength_result_headings)
document.add_paragraph(read_script_from_file(longest_dist_script_location))
populate_document_with_query_data(document, longest_dist_result, longest_dist_result_headings)


document.save('rezultati.docx')